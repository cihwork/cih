from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

REPORT_PATH = '/home/headless/workspace/cih/clientes/falqon/relatorioescuta/Relatorio_Escuta_Organizacional_Falqon_2025.docx'
LOGO_PATH = '/home/headless/workspace/cih/sobrecih/LOGO CIH azul.png'

doc = Document(REPORT_PATH)

for section in doc.sections:
    # ========== HEADER: Add logo ==========
    header = section.header
    header.is_linked_to_previous = False
    
    # Clear existing header
    for p in header.paragraphs:
        p.clear()
    
    # Add logo paragraph
    if header.paragraphs:
        header_para = header.paragraphs[0]
    else:
        header_para = header.add_paragraph()
    
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add logo
    run = header_para.add_run()
    run.add_picture(LOGO_PATH, width=Cm(3.5))
    
    # Add bottom border to header paragraph
    pPr = header_para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="8" w:color="07477D"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # ========== FOOTER: Remove logo, keep text with new URL ==========
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Clear all footer content
    for p in list(footer.paragraphs):
        p_element = p._element
        p_element.getparent().remove(p_element)
    
    # Add new footer with just text
    footer_para = footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add top border
    pPr2 = footer_para._element.get_or_add_pPr()
    pBdr2 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="8" w:color="07477D"/>'
        f'</w:pBdr>'
    )
    pPr2.append(pBdr2)
    
    # Add footer text with new URL
    footer_run = footer_para.add_run('Consultoria Impacto Humano | www.eleinepassos.com.br')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x07, 0x47, 0x7D)
    
    # Add page number
    footer_run2 = footer_para.add_run('  |  Página ')
    footer_run2.font.size = Pt(8)
    footer_run2.font.color.rgb = RGBColor(0x07, 0x47, 0x7D)
    
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r_elem = footer_run2._element
    r_new = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="16"/><w:color w:val="07477D"/></w:rPr></w:r>')
    r_new.append(fldChar1)
    footer_para._element.append(r_new)
    
    r_code = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="16"/><w:color w:val="07477D"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    footer_para._element.append(r_code)
    
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r_end = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="16"/><w:color w:val="07477D"/></w:rPr></w:r>')
    r_end.append(fldChar2)
    footer_para._element.append(r_end)
    
    # Set paragraph spacing
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0"/>')
    pPr2.append(spacing)

# Also fix any remaining old URL in document body
for para in doc.paragraphs:
    for run in para.runs:
        if 'consultoriaimpactohumano.com.br' in run.text:
            run.text = run.text.replace('consultoriaimpactohumano.com.br', 'eleinepassos.com.br')

doc.save(REPORT_PATH)
print("Relatório original corrigido:")
print("  - Logo movida para o cabeçalho (alinhada à esquerda)")
print("  - Linha separadora azul navy no cabeçalho")  
print("  - Rodapé: texto + URL www.eleinepassos.com.br + número da página")
print("  - Linha separadora no rodapé mantida")
